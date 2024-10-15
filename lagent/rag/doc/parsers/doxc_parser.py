import re
import pandas as pd
from docx import Document
from io import BytesIO
from collections import Counter
from lagent.rag.nlp import SimpleTokenizer


class DocxParser:
    def __init__(self):
        self.tokenizer = SimpleTokenizer()

    def parse_table(self, table):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        return self.format_table_content(pd.DataFrame(rows))

    def format_table_content(self, df):

        def categorize_text(text):
            patterns = [
                (r"^(20|19)\d{2}[年/-]\d{1,2}[月/-]\d{1,2}日?$", "Date"),
                (r"^(20|19)\d{2}年$", "Date"),
                (r"^(20|19)\d{2}[年/-]\d{1,2}月?$", "Date"),
                (r"^\d{1,2}[月/-]\d{1,2}日?$", "Date"),
                (r"^第*[一二三四1-4]季度$", "Date"),
                (r"^(20|19)\d{2}年*[一二三四1-4]季度$", "Date"),
                (r"^(20|19)\d{2}[ABCDE]$", "Date"),
                (r"^\d+[.,+%/ -]+$", "Number"),
                (r"^\d+[A-Z/\._~-]+$", "Code"),
                (r"^[A-Z][a-z' -]+$", "Text"),
                (r"^\d+[.,+-]+[A-Za-z/$￥%<>（）()' -]+$", "Mixed"),
                (r"^.{1}$", "Single")
            ]
            for pattern, label in patterns:
                if re.search(pattern, text):
                    return label
            tokens = [t for t in self.tokenizer.tokenize(text).split(" ") if len(t) > 1]
            if len(tokens) > 3:
                return "TextBlock" if len(tokens) < 12 else "LongText"
            # if len(tokens) == 1 and self.tokenizer.tag(tokens[0]) == "nr":
            #     return "NamedEntity"
            return "Other"

        if len(df) < 2:
            return []

        max_type = Counter(
            [categorize_text(str(df.iloc[i, j])) for i in range(1, len(df)) for j in range(len(df.columns))])
        predominant_type = max(max_type.items(), key=lambda x: x[1])[0]

        num_columns = len(df.columns)
        header_rows = [0]
        if predominant_type == "Number":
            for r in range(1, len(df)):
                types = Counter([categorize_text(str(df.iloc[r, j])) for j in range(num_columns)])
                current_type = max(types.items(), key=lambda x: x[1])[0]
                if current_type != predominant_type:
                    header_rows.append(r)

        formatted_lines = []
        for i in range(1, len(df)):
            if i in header_rows:
                continue
            header_indices = [r - i for r in header_rows if r < i]  # 表头的相对位置
            while len(header_indices) > 1 and header_indices[-1] - header_indices[-2] > 1:
                header_indices = header_indices[-1:]
            headers = []
            for j in range(num_columns):
                header_texts = []
                for idx in header_indices:
                    text = str(df.iloc[i + idx, j]).strip()
                    if text and text not in header_texts:
                        header_texts.append(text)
                header_str = ",".join(header_texts)
                if header_str:
                    header_str += ": "
                headers.append(header_str)
            cells = [headers[j] + str(df.iloc[i, j]) for j in range(num_columns) if str(df.iloc[i, j])]
            formatted_lines.append(";".join(cells))

        return formatted_lines if num_columns > 3 else ["\n".join(formatted_lines)]

    def parse(self, file_path, start_page=0, end_page=1000):
        doc = Document(file_path)
        current_page = 0

        parsed_content = {
            "paragraphs": [],
            "headings": [],
            "tables": [],
            "images": []
        }

        for para in doc.paragraphs:
            if current_page > end_page:
                break
            para_texts = []
            for run in para.runs:
                if start_page <= current_page < end_page and para.text.strip():
                    para_texts.append(run.text)
                if 'lastRenderPageBreak' in run._element.xml:
                    current_page += 1
            if para.style.name.startswith('Heading'):
                parsed_content["headings"].append({
                    "level": para.style.name,
                    "text": para.text
                })
            else:
                parsed_content["paragraphs"].append(para.text)

        for table in doc.tables:
            table_data = self.parse_table(table)
            parsed_content["tables"].append(table_data)

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                parsed_content["images"].append(rel.target_ref)

        return parsed_content


def print_parsed_content(parsed_content):
    print("Headings:")
    for heading in parsed_content.get("headings", []):
        print(f"{heading['level']}: {heading['text']}")

    print("\nParagraphs:")
    for para in parsed_content.get("paragraphs", []):
        print(f"-{para}")

    print("\nTables:")
    for table in parsed_content.get("tables", []):
        for row in table:
            print("".join(row))

    print("\nImages:")
    for image in parsed_content.get("images", []):
        print(image)


def main():
    file_path = 'example.docx'
    tokenizer = SimpleTokenizer()
    document_parser = DocxParser()
    parsed_content = document_parser.parse(file_path)
    print_parsed_content(parsed_content)


if __name__ == "__main__":
    main()
